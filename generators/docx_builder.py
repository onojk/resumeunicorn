from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

HEADER_SIZE = 16
SECTION_SIZE = 12
BODY_SIZE = 10.5

LABELS = {
    "summary": "Professional Summary",
    "skills": "Skills",
    "experience": "Experience",
    "projects": "Projects",
    "education": "Education",
}

def _add_heading(p, text, size, bold=False):
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

def _section(doc, title):
    p = doc.add_paragraph()
    _add_heading(p, title, SECTION_SIZE, bold=True)

def build_docx(data: dict) -> io.BytesIO:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_heading(p, data.get("name",""), HEADER_SIZE, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line = " | ".join([x for x in [
        data.get("title"), data.get("location"),
        data.get("phone"), data.get("email"), *data.get("links",[])
    ] if x])
    sub.add_run(line).font.size = Pt(BODY_SIZE)

    # Summary
    if data.get("summary"):
        _section(doc, LABELS["summary"])
        doc.add_paragraph(data["summary"])

    # Skills
    if data.get("skills"):
        _section(doc, LABELS["skills"])
        doc.add_paragraph(", ".join(data["skills"]))

    # Experience
    if data.get("experience"):
        _section(doc, LABELS["experience"])
        for r in data["experience"]:
            line = " — ".join([v for v in [r.get("title"), r.get("company"), r.get("location")] if v])
            hdr = doc.add_paragraph(); _add_heading(hdr, line, BODY_SIZE+1, bold=True)
            dates = " • ".join([v for v in [r.get("start"), r.get("end")] if v])
            if dates: doc.add_paragraph(dates)
            for b in r.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

    # Projects
    if data.get("options",{}).get("include_projects") and data.get("projects"):
        _section(doc, LABELS["projects"])
        for pr in data["projects"]:
            hdr = doc.add_paragraph(); _add_heading(hdr, pr.get("name",""), BODY_SIZE+1, bold=True)
            if pr.get("description"): doc.add_paragraph(pr["description"])

    # Education
    if data.get("options",{}).get("include_education") and data.get("education"):
        _section(doc, LABELS["education"])
        for ed in data["education"]:
            line = " — ".join([v for v in [ed.get("degree"), ed.get("school")] if v])
            hdr = doc.add_paragraph(); _add_heading(hdr, line, BODY_SIZE+1, bold=True)
            if ed.get("grad"): doc.add_paragraph(ed["grad"])

    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio
