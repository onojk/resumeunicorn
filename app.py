import os
import io
import re
import importlib
from flask import (
    Flask, render_template, request, redirect, url_for,
    session, abort, send_file, jsonify, make_response
)
from werkzeug.middleware.proxy_fix import ProxyFix
from flask_wtf import CSRFProtect
from flask_wtf.csrf import CSRFError
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.exceptions import RequestEntityTooLarge

# -----------------------------------------------------------------------------
# pydyf compatibility shim (MUST run before importing weasyprint)
# Some environments expose pydyf.PDF.__init__(self) while WeasyPrint>=53
# calls PDF(version, identifier). This shim normalizes the constructor
# and patches the symbol inside weasyprint.pdf as well.
# -----------------------------------------------------------------------------
_PATCH_WPDF = None
try:
    import inspect
    import pydyf
    sig = str(inspect.signature(pydyf.PDF.__init__))
    if sig == "(self)":
        _OldPDF = pydyf.PDF

        class _ShimPDF(_OldPDF):  # keep class name semantics
            def __init__(self, version="1.7", identifier=None):
                super().__init__()  # ignore args for old API

        pydyf.PDF = _ShimPDF

        def _patch_weasyprint_pdf():
            try:
                wpdf = importlib.import_module("weasyprint.pdf")
                wpdf.PDF = _ShimPDF
            except Exception:
                pass

        _PATCH_WPDF = _patch_weasyprint_pdf
except Exception:
    _PATCH_WPDF = None

# Import WeasyPrint AFTER the shim
from weasyprint import HTML, CSS
if _PATCH_WPDF:
    _PATCH_WPDF()

# DOCX
from docx import Document

# Forms
from forms import ResumeRequestForm

# -----------------------------------------------------------------------------
# App setup
# -----------------------------------------------------------------------------
app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

app.config.update(
    SECRET_KEY=os.environ.get("SECRET_KEY", "change-me"),
    MAX_CONTENT_LENGTH=256 * 1024,  # 256KB
    SESSION_COOKIE_DOMAIN=".resumeunicorn.com",
    SESSION_COOKIE_SECURE=True,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    WTF_CSRF_TIME_LIMIT=None,
    TEMPLATES_AUTO_RELOAD=True,
)

CSRFProtect(app)

limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=["60/minute", "800/hour"],
    storage_uri=os.environ.get("LIMITER_STORAGE_URI", "memory://"),
)

# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------
SAFE_ITEM = re.compile(r"[^A-Za-z0-9 +#.\-]")

def csv_to_list(raw: str, limit: int = 24):
    items = []
    for s in (raw or "").split(","):
        t = SAFE_ITEM.sub("", s).strip()
        if t:
            items.append(t)
    seen = set()
    out = []
    for i in items:
        k = i.lower()
        if k not in seen:
            seen.add(k)
            out.append(i)
    return out[:limit]

def clean_text(s: str, n: int):
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"[\x00-\x08\x0b-\x1f\x7f]", "", s)
    return s[:n]

def safe_filename(stem: str, ext: str):
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", (stem or "resume")).strip("_") or "resume"
    return f"{stem}.{ext}"

def render_resume_html(data: dict) -> str:
    """Render resume with template; fall back to a tiny inline template if missing."""
    try:
        return render_template("resume.html", data=data)
    except Exception:
        # minimal, safe fallback so downloads still work
        from markupsafe import escape
        name = escape(data.get("name", ""))
        role = escape(data.get("role", ""))
        location = escape(data.get("location", ""))
        email = escape(data.get("email", ""))
        phone = escape(data.get("phone", ""))
        linkedin = escape(data.get("linkedin", ""))
        github = escape(data.get("github", ""))
        portfolio = escape(data.get("portfolio", ""))
        summary = escape(data.get("summary", ""))

        skills_list = data.get("skills_list") or []
        certs_list = data.get("certifications_list") or []
        langs_list = data.get("languages_list") or []

        bullets = lambda xs: "".join(f"<li>{escape(x)}</li>" for x in xs)

        return f"""
<!doctype html>
<html>
<head>
<meta charset="utf-8">
<title>{name} — Resume</title>
<link rel="stylesheet" href="{request.url_root.rstrip('/')}/static/resume.css">
</head>
<body>
  <h1>{name}</h1>
  <div class="muted">{role}</div>
  <div class="muted">{location}</div>
  <div class="rule"></div>
  <p class="muted">{email} · {phone} · {linkedin} · {github} · {portfolio}</p>
  {"<h2>Summary</h2><p>"+summary+"</p>" if summary else ""}
  {"<h2>Skills</h2><ul>"+bullets(skills_list)+"</ul>" if skills_list else ""}
  {"<h2>Certifications</h2><ul>"+bullets(certs_list)+"</ul>" if certs_list else ""}
  {"<h2>Languages</h2><ul>"+bullets(langs_list)+"</ul>" if langs_list else ""}
</body>
</html>
        """

# -----------------------------------------------------------------------------
# Routes
# -----------------------------------------------------------------------------
@app.get("/healthz")
@limiter.exempt
def healthz():
    return jsonify(status="ok"), 200

@app.route("/", methods=["GET", "POST"])
@limiter.limit("10/minute; 200/hour")
def index():
    form = ResumeRequestForm()
    if request.method == "POST":
        # Honeypot: bots fill this
        if request.form.get("hp"):
            abort(400, description="Invalid input.")

        if form.validate_on_submit():
            # Normalize all fields we use later (never trust template presence)
            data = dict(
                name=clean_text(form.name.data, 80),
                email=clean_text(form.email.data, 120),
                phone=clean_text(getattr(form, "phone", type("F", (), {"data": ""}))().data, 40) if hasattr(form, "phone") else "",
                role=clean_text(form.role.data, 100),
                location=clean_text(getattr(form, "location", type("F", (), {"data": ""}))().data, 120) if hasattr(form, "location") else "",
                years=clean_text(getattr(form, "years", type("F", (), {"data": ""}))().data or "", 10) if hasattr(form, "years") else "",
                summary=clean_text(form.summary.data or "", 1200),
                linkedin=clean_text(getattr(form, "linkedin", type("F", (), {"data": ""}))().data or "", 200) if hasattr(form, "linkedin") else "",
                github=clean_text(getattr(form, "github", type("F", (), {"data": ""}))().data or "", 200) if hasattr(form, "github") else "",
                portfolio=clean_text(getattr(form, "portfolio", type("F", (), {"data": ""}))().data or "", 200) if hasattr(form, "portfolio") else "",
                work_mode=(getattr(form, "work_mode", type("F", (), {"data": ""}))().data or ""),
                theme=(getattr(form, "theme", type("F", (), {"data": ""}))().data or "emerald"),
            )

            # CSV-ish lists -> arrays
            data["skills_list"] = csv_to_list(getattr(form, "skills", type("F", (), {"data": ""}))().data)
            data["certifications_list"] = csv_to_list(
                getattr(form, "certifications", type("F", (), {"data": ""}))().data
            )
            data["languages_list"] = csv_to_list(
                getattr(form, "languages", type("F", (), {"data": ""}))().data
            )

            # simple required checks
            if not data["name"] or not data["role"] or ("@" not in data["email"]):
                abort(400, description="Invalid input.")

            # Save to session for preview & downloads (PRG)
            session["resume_data"] = data
            session.modified = True
            return redirect(url_for("success"))
        else:
            # fall through to show validation errors
            pass

    return render_template("form.html", form=form), (400 if form.errors else 200)

@app.get("/success")
def success():
    return render_template("success.html", has_data=("resume_data" in session))

@app.get("/resume")
def resume():
    data = session.get("resume_data")
    if not data:
        return redirect(url_for("index"))
    html = render_resume_html(data)
    return html

@app.get("/resume.pdf")
def resume_pdf():
    data = session.get("resume_data")
    if not data:
        abort(400, "No resume in session; please submit the form.")
    try:
        html = render_resume_html(data)
        # local CSS if present
        styles = None
        css_path = os.path.join(app.static_folder or "static", "resume.css")
        if os.path.exists(css_path):
            styles = [CSS(css_path)]
        pdf_bytes = HTML(string=html, base_url=app.root_path).write_pdf(stylesheets=styles)
        filename = safe_filename(data.get("name") or "resume", "pdf")
        resp = make_response(pdf_bytes)
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
        resp.headers["Cache-Control"] = "no-store"
        return resp
    except Exception as e:
        return make_response(f"PDF render error: {type(e).__name__}: {e}", 500)

@app.get("/resume.jpg")
def resume_jpg():
    """Render HTML → PNG with WeasyPrint, convert to JPEG via Pillow."""
    data = session.get("resume_data")
    if not data:
        abort(400, "No resume in session; please submit the form.")
    try:
        from PIL import Image
        html = render_resume_html(data)
        png_io = io.BytesIO()
        # write_png is available in modern WeasyPrint (61.x)
        HTML(string=html, base_url=app.root_path).write_png(png_io)
        png_io.seek(0)
        img = Image.open(png_io).convert("RGB")
        jpg_io = io.BytesIO()
        img.save(jpg_io, format="JPEG", quality=92, optimize=True)
        jpg_io.seek(0)
        filename = safe_filename(data.get("name") or "resume", "jpg")
        return send_file(jpg_io, mimetype="image/jpeg", as_attachment=True, download_name=filename)
    except Exception as e:
        return make_response(f"JPG render error: {type(e).__name__}: {e}", 500)

@app.get("/resume.docx")
def resume_docx():
    data = session.get("resume_data")
    if not data:
        abort(400, "No resume in session; please submit the form.")
    doc = Document()
    # Title & heading
    doc.add_heading(data.get("name", ""), level=0)
    sub_bits = [data.get("role", ""), data.get("location", ""), data.get("email", ""), data.get("phone", "")]
    doc.add_paragraph(" • ".join([x for x in sub_bits if x]))
    # Summary
    if data.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(data["summary"])
    # Skills / Certs / Languages
    for key, title in [
        ("skills_list", "Core Skills"),
        ("certifications_list", "Certifications"),
        ("languages_list", "Languages"),
    ]:
        vals = data.get(key)
        if vals:
            doc.add_heading(title, level=1)
            doc.add_paragraph(", ".join(vals))
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = safe_filename(data.get("name") or "resume", "docx")
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )

# -----------------------------------------------------------------------------
# Errors
# -----------------------------------------------------------------------------
@app.errorhandler(RequestEntityTooLarge)
def handle_413(e):
    return "Request too large", 413

@app.errorhandler(CSRFError)
def handle_csrf(e):
    ct = (request.mimetype or "")
    if request.method == "POST" and ct.startswith("application/json"):
        return "Unsupported Media Type", 415
    return (getattr(e, "description", "Bad Request")), 400

# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    app.run(host="127.0.0.1", port=8001, debug=False)
